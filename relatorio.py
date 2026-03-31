"""
Gerador de relatórios DOCX e HTML — suporta ambos os modos:
  1. Análise contra Minuta Padrão
  2. Análise de Termos Aditivos
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# =========================================================================
# Helpers
# =========================================================================
def _cor_status(status: str) -> RGBColor:
    s = status.upper()
    if s in ("OK", "CONFORME", "MANTIDA", "INFORMADO", "VALIDADO"):
        return RGBColor(0x1A, 0x7A, 0x47)
    if s in ("DIVERGENTE", "AUSENTE", "COM DIVERGÊNCIAS CRÍTICAS", "REMOVIDA"):
        return RGBColor(0xA4, 0x1C, 0x2A)
    if s in ("ALTERADO", "ALTERADA", "PARCIALMENTE CONFORME", "COM DIVERGÊNCIAS"):
        return RGBColor(0x92, 0x5C, 0x0A)
    if s == "NOVA":
        return RGBColor(0x5C, 0x3B, 0x9E)
    return RGBColor(0x6C, 0x75, 0x7D)


def _icon_status(status: str) -> str:
    s = status.upper()
    if s in ("OK", "CONFORME", "MANTIDA", "INFORMADO"):
        return "✅"
    if s in ("DIVERGENTE", "AUSENTE", "REMOVIDA"):
        return "❌"
    if s in ("ALTERADO", "ALTERADA", "PARCIALMENTE CONFORME"):
        return "🟠"
    if s == "NOVA":
        return "🟣"
    return "⚪"


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_header_row(table, headers: list):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _set_cell_shading(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)


def _add_data_row(table, values: list, bold_last=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)
        if bold_last and i == len(values) - 1:
            run.bold = True
            status = str(v).upper()
            run.font.color.rgb = _cor_status(status)


def _add_revisao_row(table, revisao: dict):
    """Adiciona linha de decisão do analista se existir."""
    if not revisao:
        return
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[-1])
    p = merged.paragraphs[0]

    status_text = "✅ VALIDADO MANUALMENTE" if revisao.get("status") == "VALIDADO" else "📋 SEGUIDO SEM VALIDAÇÃO"
    cor = "E8F5E9" if revisao.get("status") == "VALIDADO" else "FFF3E0"
    _set_cell_shading(merged, cor)

    run = p.add_run(f"Decisão: {status_text}\n")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(f"Analista: {revisao.get('analista', '—')} | {revisao.get('justificativa', '—')}")
    run.font.size = Pt(8)


# =========================================================================
# DOCX — Minuta Padrão
# =========================================================================
def gerar_relatorio_docx(resultado: dict, revisoes: dict = None) -> bytes:
    revisoes = revisoes or {}
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE ANÁLISE JURÍDICA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Conformidade com a Minuta Padrão CINCATARINA — Gestão de Frota/Combustíveis")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 1. Identificação
    doc.add_heading("1. Identificação do Contrato", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_header_row(t, ["Campo", "Valor"])
    _add_data_row(t, ["Município", resultado.get("municipio", "—")])
    _add_data_row(t, ["Número do Contrato", resultado.get("numero_contrato", "—")])
    _add_data_row(t, ["Data da Análise", resultado.get("data_analise", "—")])
    status = resultado.get("status_geral", "—")
    row = t.add_row()
    row.cells[0].paragraphs[0].add_run("Status Geral").font.size = Pt(9)
    run = row.cells[1].paragraphs[0].add_run(status)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _cor_status(status)

    # 2. Resumo
    doc.add_heading("2. Resumo Executivo", level=2)
    doc.add_paragraph(resultado.get("resumo_executivo", "—"))

    # 3. Mapeamento cláusula a cláusula
    mapeamento = resultado.get("mapeamento_clausulas", [])
    if mapeamento:
        doc.add_heading("3. Mapeamento de Cláusulas", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_header_row(t, ["Cláusula", "Contrato", "Minuta Padrão", "Status"])
        for cl in mapeamento:
            _add_data_row(t, [
                cl.get("clausula", "—"),
                cl.get("resumo_contrato", "—"),
                cl.get("resumo_minuta", "—"),
                f'{_icon_status(cl.get("status", ""))} {cl.get("status", "—")}',
            ], bold_last=True)
            obs = cl.get("observacao", "")
            if obs and obs != "—":
                row = t.add_row()
                merged = row.cells[0].merge(row.cells[3])
                run = merged.paragraphs[0].add_run(f"   ↳ {obs}")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True

    # 4. Apontamentos
    apontamentos = resultado.get("apontamentos", [])
    doc.add_heading(f"{'4' if mapeamento else '3'}. Apontamentos ({len(apontamentos)})", level=2)
    if apontamentos:
        for i, ap in enumerate(apontamentos):
            p = doc.add_paragraph()
            run = p.add_run(f"Apontamento {i+1:02d} — [{ap.get('tipo', '')}]")
            run.bold = True
            run.font.color.rgb = _cor_status(ap.get("tipo", ""))

            run2 = p.add_run(f" | {ap.get('clausula', '—')}")
            run2.font.size = Pt(10)

            t = doc.add_table(rows=1, cols=2)
            _add_header_row(t, ["Campo", "Detalhe"])
            _add_data_row(t, ["Descrição", ap.get("descricao", "—")])
            _add_data_row(t, ["Previsto na Minuta", ap.get("previsto_minuta", "—")])
            _add_data_row(t, ["Encontrado no Contrato", ap.get("encontrado_contrato", "—")])
            evidencia = ap.get("evidencia_textual", "")
            if evidencia:
                row = t.add_row()
                _set_cell_shading(row.cells[0], "EEF2FF")
                _set_cell_shading(row.cells[1], "EEF2FF")
                row.cells[0].paragraphs[0].add_run("📎 Trecho-fonte").font.size = Pt(9)
                run_ev = row.cells[1].paragraphs[0].add_run(f'"{evidencia}"')
                run_ev.font.size = Pt(8)
                run_ev.italic = True

            chave = f"apontamento_{i}"
            if chave in revisoes:
                _add_revisao_row(t, revisoes[chave])
    else:
        doc.add_paragraph("Nenhum apontamento identificado.")

    # 5. Prazos
    n = 5 if mapeamento else 4
    doc.add_heading(f"{n}. Verificação de Prazos", level=2)
    prazos = resultado.get("prazos_verificados", {})
    if prazos:
        t = doc.add_table(rows=1, cols=4)
        _add_header_row(t, ["Prazo", "Previsto", "Encontrado no Contrato", "Status"])
        nomes = {"prazo_pagamento": "Prazo de Pagamento", "prazo_vigencia": "Prazo de Vigência",
                 "prazo_correcao_vicios": "Correção de Vícios", "prazo_relatorios": "Entrega de Relatórios"}
        for chave_prazo, nome in nomes.items():
            p_data = prazos.get(chave_prazo, {})
            _add_data_row(t, [nome, p_data.get("previsto", "—"), p_data.get("encontrado", "—"),
                              f'{_icon_status(p_data.get("status", ""))} {p_data.get("status", "—")}'], bold_last=True)
            rev_key = f"prazo_{chave_prazo}"
            if rev_key in revisoes:
                _add_revisao_row(t, revisoes[rev_key])

    # 6. Multas
    doc.add_heading(f"{n+1}. Verificação de Multas e Sanções", level=2)
    multas = resultado.get("multas_verificadas", {})
    if multas:
        t = doc.add_table(rows=1, cols=4)
        _add_header_row(t, ["Sanção", "Previsto", "Encontrado", "Status"])
        nomes_m = {"multa_atraso_diaria": "Multa por Atraso (diária)", "multa_inexecucao_parcial": "Multa Inexecução Parcial",
                   "multa_inexecucao_total": "Multa Inexecução Total"}
        for chave_m, nome in nomes_m.items():
            m_data = multas.get(chave_m, {})
            _add_data_row(t, [nome, m_data.get("previsto", "—"), m_data.get("encontrado", "—"),
                              f'{_icon_status(m_data.get("status", ""))} {m_data.get("status", "—")}'], bold_last=True)
            rev_key = f"multa_{chave_m}"
            if rev_key in revisoes:
                _add_revisao_row(t, revisoes[rev_key])

    # 7. Foro e Partes
    doc.add_heading(f"{n+2}. Foro e Partes", level=2)
    t = doc.add_table(rows=1, cols=3)
    _add_header_row(t, ["Item", "Encontrado", "Status"])
    foro = resultado.get("foro_verificado", {})
    _add_data_row(t, ["Foro", foro.get("encontrado", "—"),
                       f'{_icon_status(foro.get("status", ""))} {foro.get("status", "—")}'], bold_last=True)
    if "foro" in revisoes:
        _add_revisao_row(t, revisoes["foro"])

    partes = resultado.get("partes_verificadas", {})
    nomes_p = {"contratante": "Contratante (Município)", "contratada": "Contratada (MaxiFrota)",
               "interveniente_cincatarina": "Interveniente (CINCATARINA)"}
    for chave_p, nome in nomes_p.items():
        p_data = partes.get(chave_p, {})
        _add_data_row(t, [nome, p_data.get("encontrado", "—"),
                          f'{_icon_status(p_data.get("status", ""))} {p_data.get("status", "—")}'], bold_last=True)
        rev_key = f"parte_{chave_p}"
        if rev_key in revisoes:
            _add_revisao_row(t, revisoes[rev_key])

    # 8. Decisões do Analista
    if revisoes:
        doc.add_heading(f"{n+3}. Registro de Decisões do Analista", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Apontamento", "Decisão", "Analista", "Justificativa", "Data/Hora"])
        for chave_rev, rev in revisoes.items():
            lbl = rev.get("rotulo", chave_rev)
            dec = "✅ Validado" if rev.get("status") == "VALIDADO" else "📋 Sem validação"
            _add_data_row(t, [lbl, dec, rev.get("analista", "—"), rev.get("justificativa", "—"), rev.get("timestamp", "—")])

    # Rodapé
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Relatório gerado automaticamente em {datetime.now().strftime('%d/%m/%Y')} — Analisador Jurídico CINCATARINA / MaxiFrota")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================================
# DOCX — Termos Aditivos
# =========================================================================
def gerar_relatorio_aditivos_docx(resultado: dict, revisoes: dict = None) -> bytes:
    revisoes = revisoes or {}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE ANÁLISE JURÍDICA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Análise de Termos Aditivos")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Status
    status = resultado.get("status_geral", "—")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(status)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _cor_status(status)

    # Info
    doc.add_heading("1. Identificação", level=2)
    t = doc.add_table(rows=1, cols=2)
    _add_header_row(t, ["Campo", "Valor"])
    _add_data_row(t, ["Contrato", resultado.get("numero_contrato", "—")])
    _add_data_row(t, ["Objeto", resultado.get("objeto_contrato", "—")])
    _add_data_row(t, ["Contratante", resultado.get("partes", {}).get("contratante", "—")])
    _add_data_row(t, ["Contratada", resultado.get("partes", {}).get("contratada", "—")])
    _add_data_row(t, ["Data da Análise", resultado.get("data_analise", "—")])

    # Resumo
    doc.add_heading("2. Resumo Executivo", level=2)
    doc.add_paragraph(resultado.get("resumo_executivo", "—"))

    # Mapeamento de Cláusulas
    mapeamento = resultado.get("mapeamento_clausulas", [])
    if mapeamento:
        doc.add_heading("3. Mapeamento de Cláusulas", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Cláusula", "Original", "Anteriores", "Novo Termo", "Status"])
        for cl in mapeamento:
            _add_data_row(t, [
                cl.get("clausula", "—"),
                cl.get("resumo_original", "—"),
                cl.get("resumo_anteriores", "N/A"),
                cl.get("resumo_novo_termo", "—"),
                f'{_icon_status(cl.get("status", ""))} {cl.get("status", "—")}',
            ], bold_last=True)
            obs = cl.get("observacao", "")
            if obs and obs != "—":
                row = t.add_row()
                merged = row.cells[0].merge(row.cells[4])
                run = merged.paragraphs[0].add_run(f"   ↳ {obs}")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True

    # Assinantes
    assinantes = resultado.get("assinantes_novo_termo", [])
    if assinantes:
        doc.add_heading("4. Assinantes do Novo Termo", level=2)
        t = doc.add_table(rows=1, cols=4)
        _add_header_row(t, ["Nome", "Cargo", "Presente em Anteriores", "Observação"])
        for a in assinantes:
            pres = "SIM" if a.get("presente_em_anteriores") else "NÃO"
            _add_data_row(t, [a.get("nome", "—"), a.get("cargo", "—"), pres, a.get("observacao", "—")])

    # Apontamentos
    divs = resultado.get("divergencias", [])
    doc.add_heading(f"5. Apontamentos e Divergências ({len(divs)})", level=2)
    if divs:
        t = doc.add_table(rows=1, cols=7)
        _add_header_row(t, ["Tipo", "Categoria", "Cláusula", "Previsto", "Encontrado", "Descrição", "Trecho-fonte"])
        for i, ap in enumerate(divs):
            _add_data_row(t, [
                ap.get("tipo", "—"),
                ap.get("categoria", "—"),
                ap.get("clausula", "—"),
                ap.get("previsto", "—"),
                ap.get("encontrado", "—"),
                ap.get("descricao", "—"),
                ap.get("evidencia_textual", "—"),
            ])
            chave = f"apontamento_{i}"
            if chave in revisoes:
                _add_revisao_row(t, revisoes[chave])
    else:
        doc.add_paragraph("Nenhuma divergência encontrada.")

    # Prazos
    prazos = resultado.get("prazos", [])
    if prazos:
        doc.add_heading("6. Verificação de Prazos", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Prazo", "Original", "Anteriores", "Novo Termo", "Status"])
        for p_item in prazos:
            _add_data_row(t, [p_item.get("item", "—"), p_item.get("contrato_original", "—"),
                              p_item.get("termos_anteriores", "—"), p_item.get("novo_termo", "—"),
                              f'{_icon_status(p_item.get("status", ""))} {p_item.get("status", "—")}'], bold_last=True)

    # Valores
    valores = resultado.get("valores", [])
    if valores:
        doc.add_heading("7. Verificação de Valores", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Valor/Rubrica", "Original", "Anteriores", "Novo Termo", "Status"])
        for v in valores:
            _add_data_row(t, [v.get("item", "—"), v.get("contrato_original", "—"),
                              v.get("termos_anteriores", "—"), v.get("novo_termo", "—"),
                              f'{_icon_status(v.get("status", ""))} {v.get("status", "—")}'], bold_last=True)

    # Datas
    datas = resultado.get("datas", [])
    if datas:
        doc.add_heading("8. Verificação de Datas", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Data/Evento", "Original", "Anteriores", "Novo Termo", "Status"])
        for d in datas:
            _add_data_row(t, [d.get("item", "—"), d.get("contrato_original", "—"),
                              d.get("termos_anteriores", "—"), d.get("novo_termo", "—"),
                              f'{_icon_status(d.get("status", ""))} {d.get("status", "—")}'], bold_last=True)

    # Extras/Faltantes
    extras = resultado.get("clausulas_extras", [])
    faltantes = resultado.get("clausulas_faltantes", [])
    if extras or faltantes:
        doc.add_heading("9. Cláusulas Extras e Faltantes", level=2)
        if extras:
            doc.add_paragraph("Cláusulas Extras:").runs[0].bold = True
            for e in extras:
                doc.add_paragraph(f"  • {e}")
        if faltantes:
            doc.add_paragraph("Cláusulas Faltantes:").runs[0].bold = True
            for f_item in faltantes:
                doc.add_paragraph(f"  • {f_item}")

    # Decisões
    if revisoes:
        doc.add_heading("10. Registro de Decisões do Analista", level=2)
        t = doc.add_table(rows=1, cols=5)
        _add_header_row(t, ["Apontamento", "Decisão", "Analista", "Justificativa", "Data/Hora"])
        for chave_rev, rev in revisoes.items():
            lbl = rev.get("rotulo", chave_rev)
            dec = "✅ Validado" if rev.get("status") == "VALIDADO" else "📋 Sem validação"
            _add_data_row(t, [lbl, dec, rev.get("analista", "—"), rev.get("justificativa", "—"), rev.get("timestamp", "—")])

    # Rodapé
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Relatório gerado automaticamente em {datetime.now().strftime('%d/%m/%Y')} — Validador de Termos Aditivos")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================================
# HTML Report (para ambos os modos)
# =========================================================================
def gerar_relatorio_html(resultado: dict, modo: str, revisoes: dict = None) -> str:
    revisoes = revisoes or {}
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    status = resultado.get("status_geral", "—")
    cor = "#1a7a47" if status == "CONFORME" else "#a41c2a" if "CRÍTICA" in status else "#925c0a"

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head><meta charset="UTF-8"><title>Relatório de Análise Jurídica</title>
<style>
  body{{font-family:'Segoe UI',sans-serif;font-size:12pt;color:#1a1d2e;margin:40px;background:#fff}}
  h1{{font-size:18pt;color:#1F3864;text-align:center;margin-bottom:4px}}
  h2{{font-size:13pt;color:#1F3864;margin:24px 0 8px;border-bottom:2px solid #1F3864;padding-bottom:4px}}
  .sub{{text-align:center;color:#555;font-size:10pt;margin-bottom:24px}}
  .status-box{{padding:12px 20px;border-radius:8px;font-weight:700;font-size:14pt;text-align:center;margin-bottom:20px;background:{cor}20;border:2px solid {cor};color:{cor}}}
  .resumo{{background:#eef1fa;border-radius:8px;padding:14px 16px;font-size:11pt;line-height:1.7;color:#333;margin-bottom:20px}}
  table{{width:100%;border-collapse:collapse;font-size:10pt;margin-bottom:16px}}
  th{{background:#1F3864;color:white;padding:8px 10px;text-align:left;font-size:9pt;text-transform:uppercase}}
  td{{padding:8px 10px;border-bottom:1px solid #dde1ee;vertical-align:top}}
  tr:nth-child(even) td{{background:#f9fafd}}
  .footer{{text-align:center;font-size:9pt;color:#aaa;margin-top:40px;border-top:1px solid #dde;padding-top:12px}}
  @media print{{body{{margin:20px}}}}
</style></head><body>"""

    html += f'<h1>⚖️ RELATÓRIO DE ANÁLISE JURÍDICA</h1>'

    if modo == "minuta":
        html += f'<p class="sub">Conformidade com a Minuta Padrão CINCATARINA — Gerado em {now}</p>'
        html += f'<div class="status-box">{status}</div>'
        html += f'<div class="resumo">{resultado.get("resumo_executivo", "—")}</div>'

        # Mapeamento
        mapeamento = resultado.get("mapeamento_clausulas", [])
        if mapeamento:
            html += '<h2>Mapeamento de Cláusulas</h2><table><thead><tr><th>Cláusula</th><th>Contrato</th><th>Minuta Padrão</th><th>Status</th></tr></thead><tbody>'
            for cl in mapeamento:
                s = cl.get("status", "—")
                html += f'<tr><td><b>{cl.get("clausula","—")}</b></td><td>{cl.get("resumo_contrato","—")}</td><td>{cl.get("resumo_minuta","—")}</td><td><b>{_icon_status(s)} {s}</b></td></tr>'
            html += '</tbody></table>'

        # Apontamentos
        aps = resultado.get("apontamentos", [])
        if aps:
            html += f'<h2>Apontamentos ({len(aps)})</h2><table><thead><tr><th>Tipo</th><th>Cláusula</th><th>Previsto</th><th>Encontrado</th><th>Descrição</th><th>Trecho-fonte</th></tr></thead><tbody>'
            for ap in aps:
                ev = ap.get("evidencia_textual", "")
                ev_html = f'<i>"{ev}"</i>' if ev else "—"
                html += f'<tr><td><b>{ap.get("tipo","—")}</b></td><td>{ap.get("clausula","—")}</td><td>{ap.get("previsto_minuta","—")}</td><td>{ap.get("encontrado_contrato","—")}</td><td>{ap.get("descricao","—")}</td><td style="background:#f0f4ff">{ev_html}</td></tr>'
            html += '</tbody></table>'

    else:  # termos aditivos
        html += f'<p class="sub">Análise de Termos Aditivos — Gerado em {now}</p>'
        html += f'<div class="status-box">{status}</div>'
        html += f'<div class="resumo">{resultado.get("resumo_executivo", "—")}</div>'

        mapeamento = resultado.get("mapeamento_clausulas", [])
        if mapeamento:
            html += '<h2>Mapeamento de Cláusulas</h2><table><thead><tr><th>Cláusula</th><th>Original</th><th>Anteriores</th><th>Novo Termo</th><th>Status</th></tr></thead><tbody>'
            for cl in mapeamento:
                s = cl.get("status", "—")
                html += f'<tr><td><b>{cl.get("clausula","—")}</b></td><td>{cl.get("resumo_original","—")}</td><td>{cl.get("resumo_anteriores","N/A")}</td><td>{cl.get("resumo_novo_termo","—")}</td><td><b>{_icon_status(s)} {s}</b></td></tr>'
            html += '</tbody></table>'

        divs = resultado.get("divergencias", [])
        if divs:
            html += f'<h2>Apontamentos ({len(divs)})</h2><table><thead><tr><th>Tipo</th><th>Categoria</th><th>Cláusula</th><th>Previsto</th><th>Encontrado</th><th>Descrição</th><th>Trecho-fonte</th></tr></thead><tbody>'
            for ap in divs:
                ev = ap.get("evidencia_textual", "")
                ev_html = f'<i>"{ev}"</i>' if ev else "—"
                html += f'<tr><td><b>{ap.get("tipo","—")}</b></td><td>{ap.get("categoria","—")}</td><td>{ap.get("clausula","—")}</td><td>{ap.get("previsto","—")}</td><td>{ap.get("encontrado","—")}</td><td>{ap.get("descricao","—")}</td><td style="background:#f0f4ff">{ev_html}</td></tr>'
            html += '</tbody></table>'

    # Revisões
    if revisoes:
        html += '<h2>Decisões do Analista</h2><table><thead><tr><th>Apontamento</th><th>Decisão</th><th>Analista</th><th>Justificativa</th><th>Data/Hora</th></tr></thead><tbody>'
        for k, v in revisoes.items():
            dec = "✅ Validado" if v.get("status") == "VALIDADO" else "📋 Sem validação"
            html += f'<tr><td>{v.get("rotulo", k)}</td><td>{dec}</td><td>{v.get("analista","—")}</td><td>{v.get("justificativa","—")}</td><td>{v.get("timestamp","—")}</td></tr>'
        html += '</tbody></table>'

    html += f'<div class="footer">Relatório gerado automaticamente em {now} — Analisador Jurídico CINCATARINA</div></body></html>'
    return html
